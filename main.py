from docx import Document
from faker import Faker
from presidio_analyzer import AnalyzerEngine
import re
fake = Faker()
analyzer = AnalyzerEngine()
# ----------------------------
# Email Detection
# ----------------------------
def find_emails(document):
    email_pattern = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"

    emails = set()

    for para in document.paragraphs:
        emails.update(re.findall(email_pattern, para.text))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                emails.update(re.findall(email_pattern, cell.text))

    return emails


# ----------------------------
# Phone Detection
# ----------------------------
def find_phones(document):

    phone_pattern = r"""
    (?:
        \+91[\s-]?\d{2,5}[\s-]?\d{4,8}
        |
        \b[6-9]\d{9}\b
    )
    """

    phones = set()

    # Paragraphs
    for para in document.paragraphs:

        matches = re.findall(phone_pattern, para.text, re.VERBOSE)

        for phone in matches:

            digits = re.sub(r"\D", "", phone)

            if len(digits) < 10:
                continue

            if digits.startswith(("2022", "2023", "2024")):
                continue

            phones.add(" ".join(phone.split()))

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                matches = re.findall(phone_pattern, cell.text, re.VERBOSE)

                for phone in matches:

                    digits = re.sub(r"\D", "", phone)

                    if len(digits) < 10:
                        continue

                    if digits.startswith(("2022", "2023", "2024")):
                        continue

                    phones.add(" ".join(phone.split()))

    return phones


# ----------------------------
# Email Mapping
# ----------------------------
def create_email_mapping(emails):

    mapping = {}

    for email in emails:
        mapping[email] = fake.email()

    return mapping


# ----------------------------
# Phone Mapping
# ----------------------------
# ----------------------------
# Fake Indian Phone Generator
# ----------------------------
def fake_indian_phone():
    first_digit = fake.random_element(elements=("6", "7", "8", "9"))
    remaining = "".join(str(fake.random_digit()) for _ in range(9))
    return f"+91 {first_digit}{remaining}"

# ----------------------------
# Phone Mapping
# ----------------------------
def create_phone_mapping(phones):
    mapping = {}

    for phone in phones:
        mapping[phone] = fake_indian_phone()

    return mapping


# ----------------------------
# Replace Emails
# ----------------------------
def replace_emails(document, email_mapping):

    # Paragraphs
    for para in document.paragraphs:

        for original, fake_email in email_mapping.items():

            if original in para.text:
                para.text = para.text.replace(original, fake_email)

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for original, fake_email in email_mapping.items():

                    if original in cell.text:
                        cell.text = cell.text.replace(original, fake_email)


# ----------------------------
# Replace Phones
# ----------------------------
def replace_phones(document, phone_mapping):

    # Paragraphs
    for para in document.paragraphs:

        for original, fake_phone in phone_mapping.items():

            if original in para.text:
                para.text = para.text.replace(original, fake_phone)

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for original, fake_phone in phone_mapping.items():

                    if original in cell.text:
                        cell.text = cell.text.replace(original, fake_phone)

def replace_entities(document, mapping):
    # Paragraphs
    for para in document.paragraphs:

        for original, fake_value in mapping.items():

            if original in para.text:
                para.text = para.text.replace(original, fake_value)

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for original, fake_value in mapping.items():

                    if original in cell.text:
                        cell.text = cell.text.replace(original, fake_value)

def detect_pii(text):
    return analyzer.analyze(
        text=text,
        language="en"
    )
def detect_document_pii(document):

    all_entities = []

    # Paragraphs
    for para in document.paragraphs:

        if para.text.strip():

            results = analyzer.analyze(
                text=para.text,
                language="en"
            )

            for result in results:
                entity = para.text[result.start:result.end]
                all_entities.append((result.entity_type, entity))

    # Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    results = analyzer.analyze(
                        text=cell.text,
                        language="en"
                    )

                    for result in results:
                        entity = cell.text[result.start:result.end]
                        all_entities.append((result.entity_type, entity))

    return all_entities

VALID_ENTITY_TYPES = {
    "PERSON",
    "LOCATION",
    "ORGANIZATION"
}
IGNORE_VALUES = {
    "Email",
    "Fiscals",
    "₹",
    "USD",
    "SEK",
    "SCRR",
    "RoC",
    "MoA",
    "MT"
}
IGNORE_WORDS = {
    "Circular",
    "Slip",
    "Bill",
    "Cap Price",
    "PAT Margin",
    "Schedule XIII",
    "corrigenda",
    "Refund Bank",
    "Book Building Process"
}


def is_valid_entity(entity_type, value):

    if entity_type not in VALID_ENTITY_TYPES:
        return False

    if value.strip() in IGNORE_VALUES:
        return False

    if value.strip() in IGNORE_WORDS:
        return False

    if len(value.strip()) < 3:
        return False

    # Ignore values containing digits
    if any(ch.isdigit() for ch in value):
        return False

    # Ignore emails
    if "@" in value:
        return False

    return True

# =====================================================
# Main Program
# =====================================================
def redact_document(input_path, output_path):
    doc = Document(input_path)

    emails = find_emails(doc)
    phones = find_phones(doc)

    email_mapping = create_email_mapping(emails)
    phone_mapping = create_phone_mapping(phones)

    entities = detect_document_pii(doc)

    person_mapping = {}
    organization_mapping = {}
    location_mapping = {}

    for entity_type, value in entities:

        if not is_valid_entity(entity_type, value):
            continue

        if entity_type == "PERSON":
            person_mapping.setdefault(value, fake.name())

        elif entity_type == "ORGANIZATION":
            organization_mapping.setdefault(value, fake.company())

        elif entity_type == "LOCATION":
            location_mapping.setdefault(value, fake.city())

    replace_emails(doc, email_mapping)
    replace_phones(doc, phone_mapping)
    replace_entities(doc, person_mapping)
    replace_entities(doc, location_mapping)
    replace_entities(doc, organization_mapping)

    doc.save(output_path)
if __name__ == "__main__":
    redact_document(
        "input/Red Herring Prospectus.docx",
        "output/Redacted_Prospectus.docx"
    )
